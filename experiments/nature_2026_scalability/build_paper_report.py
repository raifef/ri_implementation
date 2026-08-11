"""Build the rendered paper-style HDFA-RL versus Google RL comparison report."""
from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "13243A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "253140"
MUTED = "66717E"
GRID = "D9E0E8"
PALE = "F4F7FA"
RED = "A23A3A"
GREEN = "267653"
GOLD = "A57418"
WHITE = "FFFFFF"
PURPLE = "7B4AB5"

ARM_NAME = {
    "fixed": "Fixed calibration",
    "periodic_recalibration": "Periodic recalibration",
    "full_control_detector_rl": "Google-style full-control detector RL",
    "predictive_hdfa_no_residual": "Predictive HDFA without residual RL",
    "predictive_hdfa_residual_rl": "Predictive HDFA with residual RL",
    "oracle": "Oracle-informed controller",
}


def set_font(run, name: str = "Calibri", size: float | None = None,
             color: str | None = None, bold: bool | None = None,
             italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top: int = 80, start: int = 120,
                 bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    set_font(run, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    for name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 8),
        ("Subtitle", 14, DARK_BLUE, 0, 16),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(INK)
    caption.paragraph_format.space_before = Pt(5)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.line_spacing = 1.1
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(.375)
        style.paragraph_format.first_line_indent = Inches(-.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def setup_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(.492)
        section.footer_distance = Inches(.492)
        section.different_first_page_header_footer = True


def add_running_furniture(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("HDFA-RL SUITE | COMPARATIVE SCIENTIFIC REPORT")
    set_font(run, size=8.5, color=MUTED, bold=True)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), GRID)
    borders.append(bottom)
    p_pr.append(borders)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("Nature 2026 comparison  |  ")
    set_font(run, size=9, color=MUTED)
    add_page_field(fp)


def add_para(doc: Document, text: str, *, bold_lead: str | None = None,
             italic: bool = False, alignment=None, after: float | None = None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_font(r, italic=italic)
    else:
        r = p.add_run(text)
        set_font(r, italic=italic)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_callout(doc: Document, label: str, text: str, *, fill: str = "EEF3F8",
                accent: str = BLUE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(.12)
    p.paragraph_format.right_indent = Inches(.12)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.2
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(f"{label}  ")
    set_font(r, size=10.5, color=accent, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.5, color=INK)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
              widths: Sequence[int], *, font_size: float = 8.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    mark_header(header)
    for cell, value in zip(header.cells, headers):
        shade(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_font(r, size=font_size, color=WHITE, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            if row_index % 2 == 0:
                shade(cell, PALE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_font(r, size=font_size, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, image_path: Path, number: int, caption: str,
               alt_text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.5))
    for doc_pr in run._r.xpath(".//wp:docPr"):
        doc_pr.set("descr", alt_text)
        doc_pr.set("title", f"Figure {number}")
    cp = doc.add_paragraph(style="Caption")
    cp.paragraph_format.keep_with_next = False
    r = cp.add_run(f"Figure {number} | ")
    set_font(r, size=9, color=NAVY, bold=True)
    r = cp.add_run(caption)
    set_font(r, size=9, color=INK)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def pct(value: float) -> str:
    return f"{100*value:.0f}%"


def build(report_root: Path, output: Path) -> None:
    data = json.loads((report_root / "data" / "paper-comparison-summary.json").read_text(encoding="utf-8"))
    arms = {row["arm"]: row for row in data["arm_summary"]}
    analysis = data["scalability_analysis"]
    figures = report_root / "figures"

    doc = Document()
    configure_styles(doc)
    setup_sections(doc)
    add_running_furniture(doc.sections[0])

    # Editorial cover: narrative_proposal preset with the editorial_cover header pattern.
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("COMPARATIVE QEC CONTROL STUDY")
    set_font(r, size=11, color=GOLD, bold=True)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Hierarchical Predictive Control with Residual Reinforcement Learning")
    set_font(r, size=27, color=NAVY, bold=True)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A simulator-based effectiveness, scalability and cost comparison with Google Quantum AI's detector-driven RL framework")
    set_font(r, size=14, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("Completed v5 authoritative acquisition | 1 August 2026")
    set_font(r, size=10.5, color=MUTED, bold=True)
    add_callout(
        doc,
        "PRIMARY FINDING",
        "The experiment is authoritative but rejects the current staged implementation: all five architecture-wide gates failed, principally because the HDFA plus residual-RL arm completed only 9 of 25 runs. The scalability surrogates recover the paper's steerability anchor and show near-linear software scaling, but do not reproduce the Google hardware experiment or validate size-independent convergence.",
        fill="F7EEF0", accent=RED,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Internal simulator evidence. Not a hardware-performance or real-time deployment claim.")
    set_font(r, size=9.5, color=MUTED, italic=True)

    add_page_break(doc)
    doc.add_heading("Abstract", level=1)
    add_para(doc, (
        "We evaluate a combined hierarchical discrete fluctuation autosegmentation (HDFA), probabilistic forecasting, model-predictive control (MPC), residual reinforcement learning (RL) and lifecycle-supervision architecture for quantum error correction (QEC). The comparison uses the detector-driven full-control RL strategy of Sivak et al. as the published reference and as a faithful Google-style simulator arm. The authoritative effectiveness experiment comprised five disturbance scenarios, five independent seeds and six controller arms (150 runs), with stationary Stage 0, one held-out native-QEC baseline per matched condition, synchronized disturbance onset, complete trajectory retention and circuit-level logical evaluation using Stim and PyMatching. The staged arm completed 9/25 runs; 16/25 were censored after failed out-of-distribution Stage-0 re-entry. Full-control detector RL, fixed calibration, periodic recalibration and the oracle completed 25/25. Consequently, every predeclared acceptance gate failed, despite nominal conditional advantages in exploration damage and final detector rate. A separate scalability study covered code distances 3-15, five seeds and a declared Figure-5 surrogate with 38,670 paper-equivalent parameters at distance 15. The surrogate recovered a critical steerability frequency of 0.00681 epoch^-1 versus the paper anchor of approximately 1/150, but no convergence fit met the predeclared R-squared threshold of 0.8. The actually executed d=15 suite used 449 physical qubits and 897 controls; its staged interval required a median 12.65 s and 835 MB versus 0.143 s and 53.6 MB for the full-RL arm. The results reject superiority of the current implementation while localizing two remediable bottlenecks: lifecycle/re-entry robustness and software latency. They do not disprove the architecture in principle and do not constitute a reproduction of Willow hardware results."
    ))
    add_para(doc, "Keywords: quantum error correction; reinforcement learning; predictive control; HDFA; MPC; surface code; calibration; detector events", italic=True)

    doc.add_heading("Evidence statement", level=2)
    add_callout(doc, "BOUNDARY",
                "Published hardware measurements, declared factor-graph surrogates and executed suite probes are different evidence layers. Ratios are reported only within their valid matched layer. No suite simulation value is relabelled as a Willow measurement.")

    doc.add_heading("1. Introduction", level=1)
    add_para(doc, (
        "Continuous analogue calibration is a prerequisite for long-running QEC. Sivak et al. repurpose detector events as a learning signal for a sparse, multi-objective policy-gradient agent, avoiding the need to stop the logical computation for conventional recalibration [1]. On Willow, the published system manages more than 1,000 control parameters, improves logical-error-rate (LER) stability by 2.4x under injected drift (3.5x with decoder steering), and obtains about 20% additional LER suppression after conventional calibration. The paper also reports a 130-epoch response time, a simulated critical real-time steerability frequency near 1/150 epochs, and distance-15 scaling to almost 40,000 control parameters [1,2]."
    ))
    add_para(doc, (
        "The HDFA-RL architecture tested here makes a different systems claim. It assigns explainable structured dynamics to a joint detector-likelihood and HDFA model, propagates uncertainty through forecasting and safe MPC, and restricts RL to the residual action subspace. Stage 7 supervises every action, rollback and re-entry. This decomposition could reduce exploration damage and QEC sample cost, but it introduces model, compute and lifecycle failure modes absent from a single persistent full-control policy. The central question is therefore not whether the staged arm occasionally obtains a lower observed detector rate, but whether it completes a matched, safety-valid control experiment and passes predeclared effectiveness, sample-efficiency and no-regression gates."
    ))
    add_para(doc, (
        "This report analyses the completed v5 acquisition. It combines an authoritative six-arm effectiveness study, named circuit-level logical evidence, a paper-anchored declared scalability surrogate and an actually executed software-cost probe. The protocol was designed to preserve negative evidence: threshold non-recovery is censoring, missing data is invalidity, and 90% recovery is never inferred from an exponential fit unless that fit first passes a credibility gate."
    ))

    doc.add_heading("2. Methods", level=1)
    doc.add_heading("2.1 Comparison framework", level=2)
    add_para(doc, "The study uses three non-interchangeable layers:")
    add_bullet(doc, "Published evidence: numerical anchors and qualitative claims from the Nature article and Supplementary Information [1,2].")
    add_bullet(doc, "Declared surrogate evidence: matched factor-graph and tracking simulations constrained by published equations and protocol anchors, but not the proprietary Google Figure-5 simulator.")
    add_bullet(doc, "Executed suite evidence: the repository's actual Stage 0-7 Python paths and a named Stim/PyMatching circuit-level adapter.")
    add_para(doc, (
        "The paper's custom code is proprietary [1]. Accordingly, the surrogate comparison tests structural consistency and falsifiability; it is not a numerical reproduction of the original training environment. Public source data are referenced at Zenodo, but the 7.8 GB experimental archive is not bundled in this report [3]."
    ))

    doc.add_heading("2.2 Authoritative effectiveness design", level=2)
    add_para(doc, (
        "The primary design used a five-qubit, distance-3 simulated device; five independent seeds (101-105); five disturbances (sinusoidal local drift, semi-Markov telegraph drift, OU plus step, nested local switching under common mode, and unstructured heavy-tailed OOD drift); 32 control intervals; and 512 native-QEC cycles per interval. Each logical evaluation used 4,096 circuit shots over three rounds. Stage 0 used 384 characterization shots, 512 validation cycles, target posterior standard deviation 0.035, a 0.10 QEC detector-rate limit and an experiment-wide block-validation alpha of 10^-4."
    ))
    add_para(doc, (
        "For every scenario/seed condition, Stage 0 and a 512-cycle held-out baseline occurred before disturbance activation. Arms received identical baseline observations, cloned starting states, matched disturbance realizations and synchronized onset. Simulator truth was available only through evaluation capabilities. Every interval trajectory, policy hash, candidate trajectory, authorization, lifecycle mode, logical sample and re-entry record was preserved."
    ))
    add_para(doc, "The six independently executable arms were:")
    for arm in ARM_NAME.values():
        add_bullet(doc, arm + ".")

    doc.add_heading("2.3 Endpoints and statistics", level=2)
    add_para(doc, (
        "Primary endpoints were native-QEC cycles and candidate evaluations to observed 50%, 75% and 90% recovery; area under the excess detector-event-rate (EDR) curve; worst-region recovery; exploration damage; final EDR; circuit-level logical failure probability and logical error per round; lifecycle violations; and completion status. Recovery summaries are censoring-aware. Confidence intervals cluster the five scenario replicates by independent seed where applicable. Matched statistics require both members of the predeclared pair; incomplete primary trajectories are not imputed or silently discarded."
    ))
    add_para(doc, (
        "Five gates were predeclared: at least 10x candidate efficiency to observed 90% recovery, at least 5x reduction in integrated excess EDR, at least 2x reduction in exploration damage, at least 90% of structured runs recovering 50% within one interval, and no final-rate regression beyond a 0.005 margin. Any primary lifecycle violation or failure to complete makes the corresponding acceptance claim fail even if a conditional numerical ratio appears favourable."
    ))

    doc.add_heading("2.4 Logical-performance adapter", level=2)
    add_para(doc, (
        "Logical evidence used a named rotated surface-code memory circuit generated with Stim 1.16.0 and decoded using PyMatching 2.4.0 MWPM. The mapping from normalized control mismatch to data, measurement, reset and gate noise is evaluation-only and never exposed to a controller. This supports reproducible circuit-simulation statements; it is not a real-QPU LER measurement and is not directly commensurate with the paper's d=7 Willow LER."
    ))

    doc.add_heading("2.5 Scalability and computational-cost design", level=2)
    add_para(doc, (
        "The paper-anchored surrogate spans odd distances 3-15, one, 10 and 30 control parameters per gate, 500 epochs and five seeds. It retains the paper's 50 candidates per epoch and 36,000 QEC cycles per candidate for the Google-style arm; the staged surrogate uses four candidates at the same candidate-cycle budget. Real-time steerability uses 1,000 epochs over a 13 by 13 frequency/entropy grid. The exact d=15, P=30 structural count is 38,670 parameters."
    ))
    add_para(doc, (
        "The executed pipeline probe separately runs the actual suite at distances 3-15, corresponding to 17-449 physical qubits and 33-897 implemented line-graph controls. Each distance/seed condition executes in a fresh process, uses one common stationary Stage 0 and baseline for both counterfactual arms, and runs two online intervals. Wall time is measured without Python allocation tracing; absolute and baseline-subtracted resident memory are sampled independently. Eight workers reduce experimental wall time but do not change per-interval trajectories."
    ))

    doc.add_heading("3. Results", level=1)
    doc.add_heading("3.1 Design validity and completion", level=2)
    add_para(doc, (
        "The report is authoritative: all arms were registered, every pre-disturbance baseline was stationary, baseline hashes and disturbances were matched, onsets were synchronized, controller truth access was zero and no primary metric was missing. The overall conclusion is nevertheless negative (accepted = false). Full-control detector RL, fixed calibration, periodic recalibration and the oracle completed 25/25 runs. Predictive HDFA without residual RL completed 13/25; the full staged arm completed 9/25."
    ))
    add_para(doc, (
        "All 16 staged censors resulted from Stage-0 failure during OOD recalibration. Failure signatures involved one or both regional blocks together with the independent QEC gate, with one sensitivity failure. The staged arm accumulated five lifecycle violation events and averaged 1.8 Stage-0 executions per run. This is not missing-data attrition: the safety supervisor correctly refused to continue after loss of QEC operability. It is, however, decisive negative evidence against present lifecycle robustness."
    ))
    add_figure(doc, figures / "figure1-effectiveness-and-lifecycle.png", 1,
               "Completion, censoring-aware target attainment, and lifecycle burden across all six arms. Error bars in panel A are 95% seed-cluster bootstrap intervals; panel B uses the report's censoring-aware seed-level intervals. The staged arm's 9/25 completion rate prevents an authoritative paired superiority claim.",
               "Three-panel chart showing completion fractions, recovery target attainment, and mean Stage-0 executions for six QEC controller arms.")

    doc.add_heading("3.2 Primary trajectories and informative censoring", level=2)
    add_para(doc, (
        "Observed marginal trajectories appear to favour the staged controller: its median EDR remains near 0.04 while the full-control RL median rises toward 0.17. That contrast cannot be interpreted causally. Full RL remains at risk in all 25 runs, whereas staged follow-up falls to nine runs. Because censoring is triggered by controller-specific OOD re-entry failure, it is informative and selectively removes difficult staged trajectories. Complete-case or late-horizon marginal comparison would therefore bias toward the staged arm."
    ))
    add_para(doc, (
        "The same issue explains why the authoritative report contains no evaluable matched arm-level outcome statistics: each of the 20 comparator/outcome estimands is missing five scenario-level primary pairs after the required completion rules are applied. The correct result is non-evaluable paired superiority, not a zero effect and not a positive effect inferred from survivors."
    ))
    add_figure(doc, figures / "figure2-primary-trajectories.png", 2,
               "Interval-wise median EDR and interquartile ranges for the two primary arms, with the number of runs remaining under observation. The HDFA risk set falls from 25 to 9, so late observed differences are not a matched treatment effect.",
               "Two-panel chart comparing primary-arm detector rates and the number of runs at risk over 32 intervals.")

    doc.add_heading("3.3 Logical evidence and gate decisions", level=2)
    add_para(doc, (
        f"Across all observed horizons, the staged arm's descriptive mean circuit-level logical failure probability was {arms['predictive_hdfa_residual_rl']['mean_logical_failure_probability']:.4f}, compared with {arms['full_control_detector_rl']['mean_logical_failure_probability']:.4f} for full-control RL. Mean logical error per round was {arms['predictive_hdfa_residual_rl']['mean_logical_error_per_round']:.4g} versus {arms['full_control_detector_rl']['mean_logical_error_per_round']:.4g}. These differences are not authoritative treatment effects because the staged horizons are outcome-dependent and shorter. The figure therefore marks censored runs explicitly and avoids a cross-arm significance claim."
    ))
    add_para(doc, (
        "All five acceptance gates failed. The worst matched integrated-excess-EDR ratio was 2.07x, below the 5x target. Only 35% of structured staged runs recovered 50% within one interval, below the 90% requirement. Nominal 53.3x 90%-target sample efficiency, 31.2x exploration-damage reduction and a -0.1386 final-EDR difference were conditional on incomplete or very small complete pairs and therefore failed under the lifecycle/completion rule. This protocol behaviour is scientifically important: favourable conditional values were not promoted into an accepted result."
    ))
    add_figure(doc, figures / "figure3-logical-and-acceptance.png", 3,
               "Circuit-level logical failure probabilities and the five predeclared acceptance decisions. Hollow points are censored runs. All gates fail because the primary staged arm violates the completion/lifecycle prerequisite, in addition to shortfalls in integrated excess EDR and recurring recovery.",
               "Logical failure probability dot plot and a table showing all five architecture-wide acceptance gates failed.")

    doc.add_heading("3.4 Comparison with published Google results", level=2)
    add_para(doc, (
        "The suite reproduces one useful paper anchor at the declared-surrogate level. The Google-style surrogate's critical frequency for at least 2% mean improvement is 0.00681 epoch^-1, close to the published approximately 1/150 = 0.00667 epoch^-1. The HDFA surrogate extends to 0.0100 epoch^-1. This is evidence that the reduced tracking model is numerically anchored; it is not evidence that HDFA would extend Willow's hardware steerability boundary."
    ))
    add_para(doc, (
        "The structural sweep reaches d=15 and 38,670 paper-equivalent parameters. Conditional on both surrogate arms reaching a target, median native-QEC cycle ratios favour HDFA by 14.15x at 50% and 6.23x at 75%; there are no matched 90% completions. More importantly, none of the exponential convergence fits passes the predeclared R-squared > 0.8 criterion. Mean R-squared ranges from 0.56 to -4.60 for the Google-style surrogate and 0.24 to -45.79 for HDFA. A low coefficient of variation of gamma across distance cannot validate size-independent optimization speed when the assumed relation fits this poorly."
    ))
    add_figure(doc, figures / "figure4-google-paper-scalability-comparison.png", 4,
               "Paper anchors and declared surrogate results. The Google-style surrogate recovers the published steerability frequency, and HDFA shows conditional cycle advantages; however, 90% matched recovery is absent and every convergence fit fails the credibility threshold. These panels are not Willow hardware measurements.",
               "Four-panel comparison of steerability frequency, recovery target attainment, convergence fit quality, and native-QEC cycle ratios.")

    doc.add_heading("3.5 Executed scaling and computational cost", level=2)
    add_para(doc, (
        "All 35 distance/seed conditions and 140 method/epoch rows completed without a pipeline failure. Both implementations show approximately linear wall-time scaling in the 33-897 implemented controls: fitted exponents are 1.07 for full RL and 0.99 for HDFA. The constants differ substantially. At d=15, median interval time is 12.65 s for HDFA versus 0.143 s for full RL (about 88x), and peak process memory is 835 MB versus 53.6 MB (about 15.6x). The HDFA memory exponent is 0.77; the full-RL absolute RSS is dominated by runtime baseline and grows weakly over this range."
    ))
    add_para(doc, (
        "The paired two-epoch EDR ratio is below unity at five of seven distances, but only d=11 and d=13 exclude unity at 95%. The probe is too short for convergence or logical-scaling inference. Its computational result is clearer: the current Python staged path is not real-time at a 10-cycle, 1-ms-per-cycle interval. This does not establish hardware infeasibility. Particle filtering, scenario propagation, MPC and residual search are host-side algorithms that can be compiled, parallelized, accelerated or split across classical control hardware. It does establish that the present reference implementation cannot be deployed at the simulated cadence without substantial systems engineering."
    ))
    add_figure(doc, figures / "figure5-computational-scaling.png", 5,
               "Executed wall time, peak resident memory, relative compute cost and paired EDR ratio. Fresh processes isolate memory between conditions. At d=15 the staged reference path is about 88x slower and 15.6x larger in peak RSS than full RL; this is Python implementation evidence, not QPU latency.",
               "Four-panel computational scaling chart showing time, memory, time ratios and paired detector-rate ratios through distance 15.")

    doc.add_heading("3.6 Direct evidence-layer comparison", level=2)
    add_table(doc,
              ("Dimension", "Google RL paper", "HDFA-RL v5 evidence", "Interpretation"),
              (
                  ("Substrate", "Willow surface/color-code hardware plus proprietary simulations", "Scalable device simulator; Stim/PyMatching logical adapter", "Not a hardware reproduction"),
                  ("Control concept", "Sparse multi-objective full-control policy gradient", "Joint HDFA/dynamics + forecast + MPC + residual RL + lifecycle supervisor", "Different decomposition and failure modes"),
                  ("Experimental scale", ">1,000 controls; d=5/d=7 codes", "Effectiveness: 5-qubit d=3 simulation", "Primary effectiveness is small-scale"),
                  ("Scaling scale", "d=15; almost 40,000 simulated controls", "Declared surrogate: 38,670; executed suite: 897 controls on 449 qubits", "Structural anchor reproduced; executed control graph smaller"),
                  ("Published benefit", "20% fine-tuning LER suppression; 24% steering LER reduction; 2.4x stability", "All five acceptance gates fail; staged completion 9/25", "Current implementation does not establish relative effectiveness"),
                  ("Logical evidence", "Willow LER; d=7 surface 7.72(9)e-4 per cycle", "Circuit-simulated probability/error per round; censored primary horizons", "Metrics and substrates are not commensurate"),
                  ("Real-time/scaling", "130-epoch response; critical frequency about 1/150; size-independent gamma claim", "Anchor 0.00681 in surrogate; no credible gamma fit; d=15 interval 12.65 s", "Anchor agreement, convergence claim unvalidated, software too slow"),
              ),
              (1500, 2700, 2700, 2460), font_size=7.7)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Table 1 | Evidence-layer comparison. Published values are from Sivak et al. [1,2]; suite values are from the completed v5 artifacts.")
    set_font(r, size=8.5, color=MUTED, italic=True)

    doc.add_heading("4. Discussion", level=1)
    doc.add_heading("4.1 What the experiment proves and disproves", level=2)
    add_para(doc, (
        "The completed experiment disproves the narrow release claim that the present HDFA-RL implementation is already an effective and computationally competitive replacement for the Google-style detector RL baseline. The result is not ambiguous: the benchmark is authoritative, no design-invalidity reason remains, the full-control comparator completes, and the staged controller does not. All five gates fail."
    ))
    add_para(doc, (
        "The experiment does not disprove predictive decomposition as an architecture. Conditional runs show low exploration damage and suggest that model-based cancellation can reduce detector and logical-error proxies. The oracle and periodic recalibration arms also demonstrate that the simulated disturbances are controllable. The dominant failure is lifecycle coupling: OOD evidence triggers scientifically appropriate re-entry, but Stage 0 cannot consistently re-establish QEC operability while the disturbance regime is active. This is a specific engineering and inference problem, not a generic impossibility theorem."
    ))

    doc.add_heading("4.2 Comparison with the Google paper", level=2)
    add_para(doc, (
        "Google's published evidence remains stronger on real-device effectiveness. It measures Willow logical performance, reports repeated experimental fine-tuning and steering gains, and demonstrates a persistent policy that survives the tested drift. The suite offers broader explicit uncertainty, residual-action and lifecycle machinery, but this additional structure has not yet translated into a complete matched primary experiment. On current evidence, the Google approach is the more effective and operationally mature controller."
    ))
    add_para(doc, (
        "The suite's strongest comparative signal is sample economy within its declared surrogate: a four-candidate staged epoch can achieve large conditional native-QEC savings relative to the paper's 50-candidate protocol. Yet the 90% target is never jointly reached, and poor fit quality invalidates a size-independent convergence conclusion. The correct scientific interpretation is a promising low- and mid-recovery efficiency hypothesis requiring a more credible dynamical model and longer completed trajectories."
    ))

    doc.add_heading("4.3 Computational feasibility", level=2)
    add_para(doc, (
        "Near-linear scaling is necessary but not sufficient for real-time control. At d=15 the staged constant factor dominates: 12.65 s is approximately three orders of magnitude above the 10-ms simulated online interval. Increasing worker count shortens the offline experiment but does not reduce a single controller interval. Production feasibility therefore requires within-condition acceleration: compiled/vectorized particle likelihoods, accelerator-resident scenario propagation, parallel regional MPC, cached factorizations, bounded asynchronous forecasting and a deterministic fast safety kernel. Scientific budgets, particle counts, scenario counts and safety checks should remain unchanged while these kernels are replaced by equivalent implementations."
    ))
    add_para(doc, (
        "Stage 0 is no longer the dominant wall-time path after vectorized acquisition and graph-coloured sensitivity batching. The online predictive stack is. Any future real-time claim should report deadline miss rate, p95/p99 latency, accelerator utilization, memory transfer, supervisor latency and control-update acknowledgement on the target classical hardware, not only total benchmark wall time."
    ))

    doc.add_heading("4.4 Priority corrective experiments", level=2)
    add_bullet(doc, "Make OOD re-entry disturbance-aware: test safe hold, regional isolation and rollback before full Stage-0 reacquisition, while preserving the independent QEC-operability gate.")
    add_bullet(doc, "Run a predeclared lifecycle-focused factorial study separating Stage-2 OOD thresholds, re-entry trigger hysteresis, active-disturbance family and Stage-0 block/sensitivity failure.")
    add_bullet(doc, "Repeat the primary 25 matched conditions only after a development cohort demonstrates complete re-entry; do not tune on the held-out acceptance seeds.")
    add_bullet(doc, "Replace the origin-constrained exponential convergence gate with observed quantile times and censoring-aware survival scaling unless R-squared, residual autocorrelation and gamma uncertainty pass credibility thresholds.")
    add_bullet(doc, "Accelerate the online Stage 2-6 kernels with bit- or tolerance-equivalence tests against the reference implementation; retain full trajectory and resource accounting.")
    add_bullet(doc, "Extend circuit-level logical evidence to longer horizons and multiple code distances after primary lifecycle completion, then test detector improvement to logical improvement as a matched estimand.")

    doc.add_heading("5. Limitations", level=1)
    add_para(doc, (
        "First, the authoritative effectiveness experiment uses a five-qubit simulated device and five independent seeds. Scenario multiplicity improves stress coverage but does not create 25 independent hardware realizations; uncertainty therefore clusters by seed. Second, the Google paper's device, pulse compiler, proprietary Figure-5 simulator and training code are unavailable, preventing a like-for-like reproduction. Third, the declared factor-graph surrogate contains paper equations and anchors but not the original simulator. Fourth, the executed distance-15 suite has 897 sparse line-graph controls, not 38,670 executed physical control parameters. Fifth, two pipeline epochs measure software throughput but not convergence, long-term lifecycle stability or logical scaling. Sixth, Stim/PyMatching evidence is circuit simulation under a declared noise mapping, not hardware certification. Finally, censoring is informative in the staged arm; descriptive all-run detector and logical summaries must not be interpreted as unbiased causal effects."
    ))

    doc.add_heading("6. Conclusion", level=1)
    add_para(doc, (
        "The completed v5 data acquisition provides a decisive and scientifically useful negative result. The experiment is valid, reproducible and fully evaluable, but the current HDFA plus residual-RL product path does not meet its effectiveness claim: it completes 36% of primary runs, fails all five gates and is about 88x slower than the Google-style full-control RL reference at the largest executed distance. The suite nevertheless recovers the paper's real-time steerability anchor in a declared surrogate, demonstrates conditional native-QEC efficiency and scales its reference kernels approximately linearly in implemented controls. Those findings justify continued development, not a superiority claim. The next milestone should be a complete matched lifecycle-valid experiment, followed by accelerated online kernels and longer circuit-level logical testing. Until then, the Google paper provides the stronger evidence for effective real-device detector-driven control."
    ))

    doc.add_heading("Data and code availability", level=1)
    add_para(doc, (
        "The complete effectiveness report, matched trajectories, scalability tables, condition checkpoints, generated analysis tables and figure files are retained under artifacts/comparison/nature-2026-v5. The suite report hashes are recorded below. Google paper data are available through Zenodo [3]; the custom Google code is proprietary [1]."
    ))
    add_table(doc, ("Artifact", "Identifier"), (
        ("Effectiveness report", data["effectiveness_report_hash"]),
        ("Scalability report", data["scalability_report_hash"]),
        ("Configuration hash", data["provenance"]["configuration_hash"]),
        ("Source tree hash", data["provenance"]["source_tree_hash"]),
        ("Runtime", f"hdfa-rl-suite {data['provenance']['package_version']}; simulator {data['provenance']['simulator_version']}; Python {data['provenance']['python_version']}"),
    ), (2100, 7260), font_size=8.1)

    doc.add_heading("References", level=1)
    references = (
        "[1] Sivak, V., Morvan, A., Broughton, M. et al. Reinforcement learning control of quantum error correction. Nature 655, 879-884 (2026). https://doi.org/10.1038/s41586-026-10759-2",
        "[2] Sivak, V. et al. Supplementary Information for Reinforcement learning control of quantum error correction (2026). https://media.springernature.com/original/springer-static/esm/art%3A10.1038%2Fs41586-026-10759-2/MediaObjects/41586_2026_10759_MOESM1_ESM.pdf",
        "[3] Google Quantum AI. Data for Reinforcement Learning Control of Quantum Error Correction. Zenodo (2026). https://doi.org/10.5281/zenodo.17566521",
        "[4] Gidney, C. Stim: a fast stabilizer circuit simulator. Quantum 5, 497 (2021). https://doi.org/10.22331/q-2021-07-06-497",
        "[5] Higgott, O. PyMatching: a Python package for decoding quantum codes with minimum-weight perfect matching. ACM Transactions on Quantum Computing 3, 16 (2022). https://doi.org/10.1145/3505637",
    )
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(.25)
        p.paragraph_format.first_line_indent = Inches(-.25)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(ref)
        set_font(r, size=9.3, color=INK)

    doc.add_heading("Appendix A. Arm-level descriptive summary", level=1)
    add_para(doc, (
        "Table A1 reports all-run descriptive means. Predictive-arm values combine completed and censored horizons and must be interpreted with the completion columns; they are not matched causal estimates."
    ))
    arm_rows = []
    for row in data["arm_summary"]:
        arm_rows.append((
            ARM_NAME[row["arm"]],
            f"{row['completed']}/{row['runs']}",
            f"{row['lifecycle_violations']}",
            f"{row['mean_qec_cycles']:.0f}",
            f"{row['mean_candidate_evaluations']:.1f}",
            f"{row['mean_detector_event_rate']:.4f}",
            f"{row['mean_logical_failure_probability']:.4f}",
            f"{row['mean_exploration_damage']:.3f}",
        ))
    add_table(doc,
              ("Arm", "Complete", "Lifecycle", "QEC cycles", "Candidates", "Mean EDR", "Logical p", "Damage"),
              arm_rows, (2260, 850, 850, 980, 1050, 1000, 1200, 1170), font_size=7.2)

    doc.add_heading("Appendix B. Acceptance gate record", level=1)
    gate_rows = []
    for gate in data["acceptance_gates"]:
        ci = gate.get("confidence_interval")
        ci_text = "-" if not ci else f"[{ci['lower']:.4g}, {ci['upper']:.4g}]"
        gate_rows.append((gate["gate_id"], gate["status"].upper(),
                          f"{gate['measured_ratio']:.4g}", f"{gate['required_ratio']:.4g}",
                          str(gate["pair_count"]), ci_text))
    add_table(doc, ("Gate", "Status", "Measured", "Required", "Pairs", "95% CI"),
              gate_rows, (3430, 900, 1150, 1150, 900, 1830), font_size=7.7)
    add_para(doc, "Authoritative = true; accepted = false; invalidity reasons = none.", bold_lead="Authoritative = true;")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--report-root", type=Path,
                        default=Path("artifacts/comparison/nature-2026-v5/paper-report"))
    parser.add_argument("--output", type=Path,
                        default=Path("artifacts/comparison/nature-2026-v5/paper-report/HDFA-RL_vs_Google_RL_full_report.docx"))
    args = parser.parse_args()
    build(args.report_root, args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
